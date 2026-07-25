"""Generate ACCORD Osun proposal with detailed convincing ₦38m tech breakdown."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[2] / "docs"
OUT.mkdir(parents=True, exist_ok=True)

GREEN = RGBColor(0x14, 0x53, 0x2D)

# Detailed technology line items totaling 38,000,000
DEV_ITEMS = [
    ("Requirements, architecture & Osun deployment design", 1_500_000),
    ("Agent portal (register, claim airtime/data, unit management)", 2_500_000),
    ("Admin/command dashboard (feeds, agents, evidence, settings)", 3_000_000),
    ("Osun geo catalog (30 LGAs, 332 wards, polling-unit mapping)", 2_000_000),
    ("Live relay / streaming client integration", 2_000_000),
    ("Evidence module (snapshots, recordings, playback, download)", 1_800_000),
    ("Airtime/data claim engine + VTpass integration + anti-abuse rules", 1_700_000),
    ("ACCORD–Osun branding, UI polish & mobile responsiveness", 1_200_000),
    ("QA, security review, UAT & election-readiness testing", 1_300_000),
    ("Documentation, admin training materials & technical handover", 1_000_000),
]  # 18,000,000

HOST_ITEMS = [
    ("Production API/application hosting (election-grade uptime)", 2_500_000),
    ("Managed database hosting & high-availability configuration", 1_800_000),
    ("Media/recordings object storage & retention capacity", 2_200_000),
    ("SSL, domain, CDN & secure access layer", 800_000),
    ("Security hardening, monitoring, alerts & intrusion protection", 1_400_000),
    ("Automated backups, disaster recovery & restore drills", 1_300_000),
]  # 10,000,000

LIVE_ITEMS = [
    ("LiveKit / WebRTC realtime media infrastructure", 3_200_000),
    ("Concurrent statewide stream capacity (peak election day)", 2_800_000),
    ("Bandwidth / media relay & TURN capacity for mobile networks", 2_000_000),
    ("Stream reliability, reconnect handling & quality controls", 1_200_000),
    ("Livestream load testing, tuning & election-week surge buffer", 800_000),
]  # 10,000,000

FIELD_ITEMS = [
    ("Field agent support wallet (airtime/data)", 8_000_000),
    ("Election-day command centre operations", 4_500_000),
    ("Technical standby & field support team", 3_500_000),
    ("Devices & contingency kits (selected LGA leads)", 2_000_000),
]


def money(n: int) -> str:
    return f"{n:,}"


def set_run(run, *, bold=False, size=11, color=None, font="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", *, bold=False, size=11, align="left", space_after=6, space_before=0, color=None):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        set_run(p.add_run(text), bold=bold, size=size, color=color)
    return p


def add_h(doc, text, size=13):
    return add_para(doc, text, bold=True, size=size, align="center", space_after=6, space_before=4, color=GREEN)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = ""
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(h), bold=True, size=9)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            table.rows[r_i + 1].cells[c_i].text = ""
            is_bold = r_i == len(rows) - 1 and "TOTAL" in str(row[0]).upper()
            set_run(table.rows[r_i + 1].cells[c_i].paragraphs[0].add_run(str(val)), bold=is_bold, size=9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def section_rows(items, start_sn: int):
    rows = []
    for i, (name, amt) in enumerate(items, start=start_sn):
        rows.append([str(i), name, money(amt)])
    return rows


def build_one_pager() -> Path:
    assert sum(a for _, a in DEV_ITEMS) == 18_000_000
    assert sum(a for _, a in HOST_ITEMS) == 10_000_000
    assert sum(a for _, a in LIVE_ITEMS) == 10_000_000
    assert sum(a for _, a in FIELD_ITEMS) == 18_000_000

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.65)
        s.right_margin = Inches(0.65)

    add_para(doc, "ACCORD · OSUN STATE", bold=True, size=11, align="center", color=GREEN, space_after=1)
    add_para(doc, "ONE-PAGE EXECUTIVE BRIEF", bold=True, size=14, align="center", space_after=1)
    add_para(doc, "Technology Package — Detailed ₦38,000,000 Breakdown", bold=True, size=11, align="center", space_after=1)
    add_para(
        doc,
        "Osun State Only (30 LGAs · 332 Wards)  |  Field Operations ₦18m presented separately",
        bold=True,
        size=9,
        align="center",
        space_after=4,
    )

    add_para(doc, "Summary (Technology Package)", bold=True, size=10, color=GREEN, space_after=2)
    add_table(
        doc,
        ["Pillar", "Amount (₦)", "%"],
        [
            ["A. Platform Development & Customization", "18,000,000", "47.4%"],
            ["B. Hosting, Storage & Security", "10,000,000", "26.3%"],
            ["C. Livestreaming Infrastructure", "10,000,000", "26.3%"],
            ["TOTAL TECHNOLOGY PACKAGE", "38,000,000", "100%"],
        ],
        col_widths=[4.4, 1.4, 0.8],
    )

    add_para(doc, "A. Development — ₦18,000,000", bold=True, size=10, color=GREEN, space_before=5, space_after=2)
    add_table(
        doc,
        ["#", "Line Item", "₦"],
        section_rows(DEV_ITEMS, 1) + [["", "Subtotal — Development", "18,000,000"]],
        col_widths=[0.35, 5.3, 1.1],
    )

    add_para(doc, "B. Hosting — ₦10,000,000", bold=True, size=10, color=GREEN, space_before=5, space_after=2)
    add_table(
        doc,
        ["#", "Line Item", "₦"],
        section_rows(HOST_ITEMS, 1) + [["", "Subtotal — Hosting", "10,000,000"]],
        col_widths=[0.35, 5.3, 1.1],
    )

    add_para(doc, "C. Livestreaming — ₦10,000,000", bold=True, size=10, color=GREEN, space_before=5, space_after=2)
    add_table(
        doc,
        ["#", "Line Item", "₦"],
        section_rows(LIVE_ITEMS, 1) + [["", "Subtotal — Livestreaming", "10,000,000"]],
        col_widths=[0.35, 5.3, 1.1],
    )

    add_para(doc, "Separate Field Operations — ₦18,000,000 (not in ₦38m)", bold=True, size=10, color=GREEN, space_before=5, space_after=2)
    add_table(
        doc,
        ["Line Item", "₦"],
        [[n, money(a)] for n, a in FIELD_ITEMS] + [["TOTAL FIELD OPS (SEPARATE)", "18,000,000"]],
        col_widths=[5.5, 1.2],
    )

    add_para(
        doc,
        "Decision: Approve Technology Package ₦38,000,000. Separately consider Field Ops ₦18,000,000.",
        bold=True,
        size=9,
        space_before=5,
        space_after=3,
    )
    add_para(doc, "Approved by: Name _____________ Signature _____________ Date _________", size=9, space_after=1)

    path = OUT / "ACCORD_Osun_Executive_Brief_38m.docx"
    doc.save(path)
    return path


def build_full_proposal() -> Path:
    assert sum(a for _, a in DEV_ITEMS) == 18_000_000
    assert sum(a for _, a in HOST_ITEMS) == 10_000_000
    assert sum(a for _, a in LIVE_ITEMS) == 10_000_000

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.85)
        s.bottom_margin = Inches(0.85)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # COVER
    for _ in range(2):
        add_para(doc, "", space_after=6)
    add_para(doc, "ACCORD", bold=True, size=28, align="center", color=GREEN, space_after=4)
    add_para(doc, "OSUN STATE", bold=True, size=20, align="center", space_after=10)
    add_para(doc, "FORMAL PROPOSAL", bold=True, size=14, align="center", space_after=8)
    add_para(doc, "Election Monitoring, Command", bold=True, size=18, align="center", space_after=2)
    add_para(doc, "& Vote Protection Platform", bold=True, size=18, align="center", space_after=10)
    add_para(doc, "Technology Package (Detailed Breakdown)", bold=True, size=12, align="center", space_after=2)
    add_para(doc, "₦38,000,000", bold=True, size=22, align="center", color=GREEN, space_after=2)
    add_para(doc, "(Thirty-Eight Million Naira Only)", bold=True, size=11, align="center", space_after=8)
    add_para(doc, "Separate Field Operations Budget: ₦18,000,000", bold=True, size=11, align="center", space_after=10)
    add_para(doc, "Coverage: Osun State Only · 30 LGAs · 332 Wards", size=11, align="center", space_after=14)
    add_para(doc, "Presented to: ACCORD Leadership / Campaign Council", size=11, align="center", space_after=2)
    add_para(doc, "Prepared by: _______________________________", size=11, align="center", space_after=2)
    add_para(doc, "Date: _______________________________", size=11, align="center", space_after=2)
    add_para(doc, "Classification: Confidential", bold=True, size=11, align="center")
    doc.add_page_break()

    add_h(doc, "TABLE OF CONTENTS", 14)
    for t in [
        "1. Executive Summary",
        "2. Why a Detailed Technology Breakdown Matters",
        "3. Scope (Osun State Only)",
        "4. Technology Package Summary — ₦38,000,000",
        "5. Pillar A — Platform Development (₦18,000,000)",
        "6. Pillar B — Hosting, Storage & Security (₦10,000,000)",
        "7. Pillar C — Livestreaming Infrastructure (₦10,000,000)",
        "8. Separate Field Operations Budget — ₦18,000,000",
        "9. Implementation Timeline",
        "10. Disbursement Schedule",
        "11. Decision Request",
        "12. Signature and Approval Page",
    ]:
        add_para(doc, t, size=11, space_after=4)
    doc.add_page_break()

    add_h(doc, "1. EXECUTIVE SUMMARY")
    add_para(
        doc,
        "This proposal requests approval of ₦38,000,000 for the Technology Package of ACCORD’s "
        "Osun State Election Monitoring Platform. To support transparent leadership review, the "
        "₦38 million is broken into three clear pillars — Development, Hosting, and Livestreaming — "
        "with detailed line items under each pillar. Field operations (₦18,000,000) remain a separate budget.",
        size=11,
        align="justify",
        space_after=10,
    )

    add_h(doc, "2. WHY A DETAILED TECHNOLOGY BREAKDOWN MATTERS")
    for item in [
        "Shows exactly what ACCORD is paying for — not a vague lump sum",
        "Links every naira to a delivery that improves election-day visibility and evidence",
        "Makes milestone payments easier to verify",
        "Separates software/infrastructure from field logistics for cleaner governance",
        "Helps leadership defend the investment to stakeholders with clear value lines",
    ]:
        add_para(doc, f"• {item}", size=11, space_after=2)
    add_para(doc, "", space_after=6)

    add_h(doc, "3. SCOPE (OSUN STATE ONLY)")
    add_table(
        doc,
        ["Item", "Coverage"],
        [
            ["State", "Osun State only"],
            ["LGAs / Wards", "30 LGAs · 332 Wards"],
            ["Technology Package", "₦38,000,000 (Development + Hosting + Livestreaming)"],
            ["Field Operations", "₦18,000,000 (separate approval)"],
        ],
        col_widths=[2.3, 4.2],
    )
    add_para(doc, "", space_after=8)

    add_h(doc, "4. TECHNOLOGY PACKAGE SUMMARY — ₦38,000,000")
    add_table(
        doc,
        ["Pillar", "Amount (₦)", "Share", "What it buys"],
        [
            ["A. Platform Development & Customization", "18,000,000", "47.4%", "Build/adapt the full ACCORD–Osun system"],
            ["B. Hosting, Storage & Security", "10,000,000", "26.3%", "Keep the system online, secure, and backed up"],
            ["C. Livestreaming Infrastructure", "10,000,000", "26.3%", "Carry concurrent live feeds on election day"],
            ["TOTAL", "38,000,000", "100%", ""],
        ],
        col_widths=[2.6, 1.2, 0.8, 1.9],
    )
    add_para(doc, "", space_after=8)

    add_h(doc, "5. PILLAR A — PLATFORM DEVELOPMENT (₦18,000,000)")
    add_para(
        doc,
        "This pillar funds the actual product ACCORD will use: portals, dashboards, Osun mapping, "
        "evidence tools, claim controls, branding, testing, and handover.",
        size=11,
        align="justify",
        space_after=6,
    )
    add_table(
        doc,
        ["S/N", "Detailed Line Item", "Amount (₦)"],
        section_rows(DEV_ITEMS, 1) + [["", "SUBTOTAL — DEVELOPMENT", "18,000,000"]],
        col_widths=[0.5, 4.6, 1.4],
    )
    add_para(doc, "", space_after=6)
    add_para(doc, "Value to leadership:", bold=True, size=11, space_after=3)
    add_para(
        doc,
        "A working statewide command system — not just a website — that agents and admins can use "
        "immediately for monitoring, evidence, and controlled field support workflows.",
        size=11,
        align="justify",
        space_after=10,
    )

    add_h(doc, "6. PILLAR B — HOSTING, STORAGE & SECURITY (₦10,000,000)")
    add_para(
        doc,
        "This pillar funds the production environment that must stay available through the election "
        "window: servers, database, media storage, security, and recovery.",
        size=11,
        align="justify",
        space_after=6,
    )
    add_table(
        doc,
        ["S/N", "Detailed Line Item", "Amount (₦)"],
        section_rows(HOST_ITEMS, 1) + [["", "SUBTOTAL — HOSTING & SECURITY", "10,000,000"]],
        col_widths=[0.5, 4.6, 1.4],
    )
    add_para(doc, "", space_after=6)
    add_para(doc, "Value to leadership:", bold=True, size=11, space_after=3)
    add_para(
        doc,
        "Election-day reliability and data protection — feeds, evidence, and admin access remain "
        "available when pressure is highest.",
        size=11,
        align="justify",
        space_after=10,
    )

    add_h(doc, "7. PILLAR C — LIVESTREAMING INFRASTRUCTURE (₦10,000,000)")
    add_para(
        doc,
        "This pillar funds realtime video capacity for agents streaming from polling units across "
        "Osun, including peak concurrent load on election day and mobile-network resilience.",
        size=11,
        align="justify",
        space_after=6,
    )
    add_table(
        doc,
        ["S/N", "Detailed Line Item", "Amount (₦)"],
        section_rows(LIVE_ITEMS, 1) + [["", "SUBTOTAL — LIVESTREAMING", "10,000,000"]],
        col_widths=[0.5, 4.6, 1.4],
    )
    add_para(doc, "", space_after=6)
    add_para(doc, "Value to leadership:", bold=True, size=11, space_after=3)
    add_para(
        doc,
        "True live visibility from the field — leadership can see what is happening at polling units "
        "in near real time, not only after-the-fact reports.",
        size=11,
        align="justify",
        space_after=10,
    )

    add_h(doc, "8. SEPARATE FIELD OPERATIONS BUDGET — ₦18,000,000")
    add_para(
        doc,
        "These items are intentionally outside the ₦38 million Technology Package so leadership can "
        "approve technology first and field logistics separately.",
        size=11,
        align="justify",
        space_after=6,
    )
    add_table(
        doc,
        ["S/N", "Line Item", "Amount (₦)"],
        [[str(i), n, money(a)] for i, (n, a) in enumerate(FIELD_ITEMS, 1)]
        + [["", "TOTAL FIELD OPERATIONS (SEPARATE)", "18,000,000"]],
        col_widths=[0.5, 4.6, 1.4],
    )
    add_para(doc, "Combined if both packages approved: ₦56,000,000", bold=True, size=11, space_before=6, space_after=10)

    add_h(doc, "9. IMPLEMENTATION TIMELINE")
    add_table(
        doc,
        ["Phase", "Period", "Focus"],
        [
            ["1", "Week 1–2", "Kickoff, architecture, hosting setup, branding"],
            ["2", "Week 3–4", "Core modules, Osun catalog, pilot LGAs"],
            ["3", "Week 5–6", "Livestream capacity, security hardening, statewide readiness"],
            ["4", "Election window", "Live operations (field package if approved)"],
            ["5", "Post-election", "Evidence archive, handover, final report"],
        ],
        col_widths=[0.8, 1.4, 4.3],
    )
    add_para(doc, "", space_after=8)

    add_h(doc, "10. DISBURSEMENT SCHEDULE — TECHNOLOGY ₦38m")
    add_table(
        doc,
        ["Tranche", "%", "Amount (₦)", "Trigger"],
        [
            ["1", "40%", "15,200,000", "Approval + development kickoff"],
            ["2", "35%", "13,300,000", "Hosting live + livestream staging verified"],
            ["3", "25%", "9,500,000", "Statewide technical readiness / final delivery"],
            ["", "100%", "38,000,000", ""],
        ],
        col_widths=[1.0, 0.8, 1.5, 3.2],
    )
    add_para(doc, "", space_after=8)

    add_h(doc, "11. DECISION REQUEST")
    for i, item in enumerate(
        [
            "Approve Technology Package ₦38,000,000 with the detailed pillar breakdown above",
            "Separately consider Field Operations Package ₦18,000,000",
            "Authorize milestone-based disbursement",
            "Nominate Osun LGA focal persons",
            "Schedule leadership demonstration within 7 days of approval",
        ],
        start=1,
    ):
        add_para(doc, f"{i}. {item}", size=11, space_after=2)
    add_para(doc, "", space_after=8)

    doc.add_page_break()
    add_h(doc, "12. SIGNATURE AND APPROVAL PAGE", 14)
    add_para(doc, "A. Technology Package — ₦38,000,000", bold=True, size=12, color=GREEN, space_after=4)
    add_para(doc, "Development ₦18,000,000 · Hosting ₦10,000,000 · Livestreaming ₦10,000,000", size=10, space_after=8)
    add_para(doc, "Prepared By: Name ______________ Signature ______________ Date ________", size=10, space_after=4)
    add_para(doc, "Reviewed By: Name ______________ Signature ______________ Date ________", size=10, space_after=4)
    add_para(doc, "Approved By (ACCORD Leadership):", bold=True, size=11, space_after=4)
    add_para(doc, "Name ______________ Title ______________ Signature ______________ Date ________", size=10, space_after=4)
    add_para(doc, "Amount Approved: ₦38,000,000 (Technology Package — Osun Only)", bold=True, size=11, space_after=12)

    add_para(doc, "B. Field Operations Package — ₦18,000,000 (Separate)", bold=True, size=12, color=GREEN, space_after=4)
    add_para(doc, "Wallet ₦8m · Command Centre ₦4.5m · Standby Team ₦3.5m · Devices ₦2m", size=10, space_after=8)
    add_para(doc, "Approved By: Name ______________ Signature ______________ Date ________", size=10, space_after=4)
    add_para(doc, "Amount Approved: ₦18,000,000 (Field Operations — Separate / Optional)", bold=True, size=11, space_after=8)
    add_para(doc, "Official Stamp / Seal: _______________________________", size=11)

    path = OUT / "ACCORD_Osun_Formal_Proposal_38m_Detailed_Tech.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print("dev", sum(a for _, a in DEV_ITEMS))
    print("host", sum(a for _, a in HOST_ITEMS))
    print("live", sum(a for _, a in LIVE_ITEMS))
    print(build_one_pager())
    print(build_full_proposal())
